#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
scrie_document.py

Scrie un act complet in forma casei, pornind de la sablon si de la un continut dat.

De ce exista. Sablonul aduce definitiile de stil, dar cine scrie continutul poate ignora
tot. Documentele generate ieseau cu marginile corecte si structura altcuiva, pe `Normal`
in loc de stilul de corp, cu numerotare tastata sau lipsa, aproape fara bold.

Aici formatarea nu mai e la latitudinea celui care scrie. Textul si rolul fiecarui bloc
vin din afara, iar stilul, numerotarea si bold-ul se aplica de aici, dupa masuratoarea pe
cele 86 de acte proprii.

    python scrie_document.py --continut act.json --output cerere.docx

Formatul continutului
---------------------
{
  "titlu": "cerere de chemare in judecata",
  "instanta": ["CATRE: TRIBUNALUL ARGES", "SECTIA DE CONTENCIOS ADMINISTRATIV SI FISCAL"],
  "reclamant": ["Subsemnata, X, cu domiciliul in ...",
                "Avand domiciliul procesual ales, in temeiul art. 158 C. pr. civ., ..."],
  "parti": ["In contradictoriu cu",
            {"text": "INSPECTORATUL DE POLITIE AL JUDETULUI ARGES", "bold": true},
            "in temeiul Legii nr. 554/2004 a contenciosului administrativ, formulez"],
  "obiect": ["impotriva avizului negativ nr. 384746 din 08.12.2025"],
  "corp": [
    "Paragraf obisnuit, primeste automat numarul [1].",
    {"text": "Termenul de arma interzisa nu este definit.", "bold": ["arma interzisa"]},
    {"tip": "marcator", "text": "element de enumerare, primeste liniuta"},
    {"tip": "titlu", "text": "I. IN FAPT"}
  ],
  "final": ["Cu deosebita consideratie,", "Av. Adrian Zamfir"]
}

Orice bloc lipseste, se sare. `bold` primeste fie true pentru tot paragraful, fie o lista
de fragmente de evidentiat in text.
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("Lipseste python-docx. Instaleaza-l cu: pip install python-docx")

AICI = Path(__file__).resolve().parent
SABLON = AICI.parent / "assets" / "sablon-casa.docx"

sys.path.insert(0, str(AICI))
from cale_libera import alege, adauga_optiune
import repara_pachet  # noqa: E402  spatiile de nume se verifica dupa fiecare salvare

# Masurat pe actele proprii. Corpul se numeroteaza [1], [2], [3], cu paranteze drepte,
# prin fluxul 41; enumerarile din interiorul corpului folosesc liniuta, fluxul 37.
# numId=0 anuleaza numerotarea mostenita din stil, folosit la blocul partilor.
NUM_CORP = 41
NUM_MARCATOR = 37
NUM_FARA = 0

# Numarul de paragraf atarna in marginea din stanga, iar textul incepe chiar la margine.
# Masurat pe actele proprii: toate cele 225 de paragrafe numerotate din cererea de camera
# preliminara Transcarpat poarta exact `w:ind w:left="0" w:hanging="720"`, fara exceptie.
# Fara aceste valori, Word cade pe indentarea din definitia fluxului, left=360
# hanging=360, iar numarul se aseaza cu 1,27 cm mai la dreapta decat in actele proprii.
INDENT_CORP = (0, 720)

# Blocurile pe stilul simplu, adresarea, identificarea partilor, obiectul si finalul, se
# aliniaza la marginea din stanga, ca textul de corp. Stilul `Normal` poarta left=504, iar
# sablonul are aceeasi valoare ca implicit al documentului, deci un paragraf fara `w:ind`
# iese cu 0,89 cm mai la dreapta decat corpul. In actele proprii fiecare paragraf pe
# `Normal` scrie explicit left=0; masurat pe cererea de camera preliminara Transcarpat,
# fara nicio exceptie. Semnalat de autor pe 30 iulie 2026.
INDENT_SIMPLU = 0

# Spatierea din jurul titlurilor, in twips. Valorile sunt cele dominante in actele
# proprii: `w:before="160"` apare de 88 de ori in cererea de camera preliminara
# Transcarpat, iar `w:after="100"` de 154 de ori.
#
# Titlul principal se centreaza si primeste spatiu de o parte si de alta. Titlurile de
# sectiune primesc spatiu inainte, ca cititorul sa vada unde incepe una noua.
#
# Cand titlul cade in capul unei pagini, spatiul dinainte nu mai are rost, pagina noua
# separa singura. Word ignora deja `space before` la inceput de pagina cand ruperea e
# automata; pentru ruperile puse de mana, comutatorul `suppressSpBfAfterPgBrk` din
# settings.xml face acelasi lucru, vezi `_fara_spatiu_in_capul_paginii`.
SPATIU_TITLU_PRINCIPAL = (160, 160)
SPATIU_TITLU_SECTIUNE = (160, 100)

# Ordinea copiilor lui w:pPr din schema. Word refuza documentul cand nu e respectata,
# deci elementele noi se insereaza pe pozitie, nu la coada.
ORDINE_PPR = (
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr", "widowControl",
    "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs", "suppressAutoHyphens",
    "kinsoku", "wordWrap", "overflowPunct", "topLinePunct", "autoSpaceDE", "autoSpaceDN",
    "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind", "contextualSpacing",
    "mirrorIndents", "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr", "pPrChange",
)

# python-docx cere numele stilului, nu identificatorul. Cautarea dupa id e depreciata si
# scoate un avertisment la fiecare paragraf.
STIL_CORP = "Body cu nr de paragraf"
STIL_PARTI = "Parties"
STIL_TITLU = "heading 1"
STIL_SIMPLU = "Normal"

# Titlurile se scriu bold si albastru, pe run, nu prin stil. Masurat pe titlurile din
# cele 86 de acte proprii: bold in 100% din cazuri, culoarea vazuta de cititor #244061
# in 54%, urmata de #1F497D cu 27%. Definitia stilului `Heading1` poarta un violet
# #590056, dar el este suprascris de fiecare data, deci nu ajunge niciodata pe hartie.
TITLU_CULOARE = "244061"


def _pune_in_pPr(ppr, element, nume: str):
    """Insereaza elementul pe pozitia ceruta de schema, nu la coada."""
    for vechi in ppr.findall(qn("w:" + nume)):
        ppr.remove(vechi)
    rang = ORDINE_PPR.index(nume)
    for copil in ppr:
        eticheta = copil.tag.split("}")[-1]
        if eticheta in ORDINE_PPR and ORDINE_PPR.index(eticheta) > rang:
            copil.addprevious(element)
            return
    ppr.append(element)


def _spatiere(paragraf, inainte: int, dupa: int, centrat: bool = False):
    """Scrie spatierea si, la cerere, centrarea direct pe paragraf.

    Direct pe paragraf, nu in stil, fiindca asa arata actele proprii: definitia lui
    `Heading1` nu poarta nici spatiere, nici aliniere, totul e pus pe paragraf.
    """
    ppr = paragraf._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(inainte))
    sp.set(qn("w:after"), str(dupa))
    _pune_in_pPr(ppr, sp, "spacing")
    if centrat:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        _pune_in_pPr(ppr, jc, "jc")


def _aliniaza_la_margine(paragraf):
    """Scrie `w:ind w:left=0` pe paragraf, ca textul sa porneasca de la marginea din
    stanga, aliniat cu corpul actului. Vezi INDENT_SIMPLU pentru de ce nu ajunge stilul."""
    ppr = paragraf._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(INDENT_SIMPLU))
    _pune_in_pPr(ppr, ind, "ind")


def _fara_spatiu_in_capul_paginii(doc):
    """Spune Word sa nu aplice spatiul dinainte cand paragraful deschide o pagina.

    Word face deja asta la ruperile automate de pagina. Comutatorul acopera si ruperile
    puse de mana, ca regula sa fie una singura: titlul primeste spatiu inainte peste tot,
    in afara de capul paginii, unde pagina noua separa singura.
    """
    settings = doc.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    if compat.find(qn("w:suppressSpBfAfterPgBrk")) is not None:
        return
    flag = OxmlElement("w:suppressSpBfAfterPgBrk")
    # In CT_Compat, comutatoarele stau inaintea elementelor w:compatSetting.
    primul_setting = compat.find(qn("w:compatSetting"))
    if primul_setting is not None:
        primul_setting.addprevious(flag)
    else:
        compat.append(flag)


def _numerotare(paragraf, num_id: int, nivel: int = 0, indentare=None):
    """Ataseaza fluxul de numerotare direct pe paragraf.

    Numerele nu se scriu niciodata in text. Unul tastat se strica la prima insertie si
    apare ca modificare intr-un redline, desi cititorul vede acelasi lucru.

    `indentare` primeste (stanga, atarnare) in twips si se scrie ca `w:ind` direct pe
    paragraf. Fara ea, Word cade pe indentarea din definitia fluxului si numarul se
    aseaza altfel decat in actele proprii, vezi INDENT_CORP.
    """
    ppr = paragraf._p.get_or_add_pPr()
    for vechi in ppr.findall(qn("w:numPr")):
        ppr.remove(vechi)
    npr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(nivel))
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    npr.append(ilvl)
    npr.append(nid)
    ppr.append(npr)

    if indentare is None:
        return
    stanga, atarnare = indentare
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(stanga))
    ind.set(qn("w:hanging"), str(atarnare))
    _pune_in_pPr(ppr, ind, "ind")


def _scrie(paragraf, text: str, bold):
    """Aseaza textul, cu evidentierile cerute.

    `bold` poate fi True pentru tot paragraful sau o lista de fragmente. Fragmentele se
    cauta in ordine si textul se sparge in run-uri, ca evidentierea sa cada exact pe ele.
    """
    if bold is True:
        paragraf.add_run(text).bold = True
        return
    if not bold:
        paragraf.add_run(text)
        return

    ramas = text
    for fragment in bold:
        if not fragment or fragment not in ramas:
            continue
        inainte, _, ramas = ramas.partition(fragment)
        if inainte:
            paragraf.add_run(inainte)
        paragraf.add_run(fragment).bold = True
    if ramas:
        paragraf.add_run(ramas)


def _culoare_titlu(paragraf):
    """Pune albastrul pe run, peste ce spune stilul.

    Definitia stilului `Heading1` din sablon poarta violet, dar in toate actele proprii
    culoarea e suprascrisa pe run, deci cititorul vede mereu albastru. Reproducem ce se
    vede, nu ce scrie in definitie.
    """
    for run in paragraf.runs:
        rpr = run._r.get_or_add_rPr()
        for vechi in rpr.findall(qn("w:color")):
            rpr.remove(vechi)
        col = OxmlElement("w:color")
        col.set(qn("w:val"), TITLU_CULOARE)
        rpr.append(col)


def _bloc(doc, elemente, stil, *, bold_implicit=False, num_id=None):
    for e in elemente or []:
        item = {"text": e} if isinstance(e, str) else dict(e)
        text = item.get("text", "")
        if not text:
            continue

        tip = item.get("tip", "")
        stil_efectiv = stil
        numerotare = num_id
        e_titlu = tip == "titlu"
        if tip == "marcator":
            numerotare = NUM_MARCATOR
        elif e_titlu:
            stil_efectiv = STIL_TITLU
            numerotare = None

        try:
            p = doc.add_paragraph(style=stil_efectiv)
        except KeyError:
            p = doc.add_paragraph()

        bold = True if e_titlu else item.get("bold", bold_implicit)
        _scrie(p, text, bold)
        if e_titlu:
            _culoare_titlu(p)
            _spatiere(p, *SPATIU_TITLU_SECTIUNE)
        if numerotare is not None:
            # Numai corpul primeste indentarea masurata. Enumerarile cu liniuta stau in
            # interiorul paragrafului si isi pastreaza retragerea din definitia fluxului.
            _numerotare(p, numerotare,
                        indentare=INDENT_CORP if numerotare == NUM_CORP else None)
        elif stil_efectiv == STIL_SIMPLU:
            _aliniaza_la_margine(p)
        yield p


def construieste(spec: dict, autor: str):
    if not SABLON.exists():
        sys.exit(f"Lipseste sablonul: {SABLON}")
    doc = Document(str(SABLON))

    # Sablonul pastreaza un paragraf gol, ca documentul sa fie valid. Il scoatem dupa ce
    # avem primul continut, altfel actul incepe cu un rand liber.
    gol = doc.paragraphs[0] if doc.paragraphs else None

    numarate = 0
    # Antetul de adresare, identificarea si obiectul stau pe stilul simplu, cu bold.
    list(_bloc(doc, spec.get("instanta"), STIL_SIMPLU, bold_implicit=True))
    list(_bloc(doc, spec.get("reclamant"), STIL_SIMPLU, bold_implicit=True))
    # Blocul partilor isi anuleaza numerotarea mostenita din stil.
    list(_bloc(doc, spec.get("parti"), STIL_PARTI, num_id=NUM_FARA))

    if spec.get("titlu"):
        try:
            p = doc.add_paragraph(style=STIL_TITLU)
        except KeyError:
            p = doc.add_paragraph()
        p.add_run(spec["titlu"]).bold = True
        _culoare_titlu(p)
        _spatiere(p, *SPATIU_TITLU_PRINCIPAL, centrat=True)

    list(_bloc(doc, spec.get("obiect"), STIL_SIMPLU, bold_implicit=True))

    for p in _bloc(doc, spec.get("corp"), STIL_CORP, num_id=NUM_CORP):
        numarate += 1

    list(_bloc(doc, spec.get("final"), STIL_SIMPLU))

    if gol is not None and len(doc.paragraphs) > 1:
        gol._element.getparent().remove(gol._element)

    _fara_spatiu_in_capul_paginii(doc)

    core = doc.core_properties
    core.author = autor
    core.last_modified_by = autor
    return doc, numarate


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Act complet in forma casei")
    ap.add_argument("--continut", required=True, type=Path, help="fisier JSON cu continutul")
    ap.add_argument("--output", required=True, type=Path)
    adauga_optiune(ap)
    ap.add_argument("--autor", default="Adrian Zamfir")
    args = ap.parse_args(argv)

    spec = json.loads(args.continut.read_text(encoding="utf-8"))
    doc, numarate = construieste(spec, args.autor)

    args.output = alege(args.output, args.suprascrie)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))
    # Word refuza un pachet in care Ignorable enumera prefixe nedeclarate, chiar daca
    # fisierul e XML valid. Serializarea poate reintroduce defectul, deci se verifica
    # dupa fiecare salvare, nu doar cand sablonul e suspect.
    repara_pachet.asigura(args.output)

    cu_text = sum(1 for p in doc.paragraphs if p.text.strip())
    cu_bold = sum(1 for p in doc.paragraphs if p.text.strip() and any(r.bold for r in p.runs))
    print(f"{args.output}")
    print(f"  paragrafe cu text: {cu_text}, din care numerotate automat: {numarate}")
    print(f"  paragrafe cu bold: {cu_bold} ({100 * cu_bold // max(cu_text, 1)}%), "
          f"in actele proprii 53%")
    print(f"  verifica cu: python ../../docx-livrare-check/scripts/check_livrare.py "
          f"\"{args.output}\"")
    return 0


if __name__ == "__main__":
    sys.exit(main())
