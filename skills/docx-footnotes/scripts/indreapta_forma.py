#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
indreapta_forma.py

Aduce in forma casei paragrafele adaugate intr-un act care era deja in forma casei.

De ce exista
------------
Regulile de forma acopereau scrierea actului din sablon, prin `scrie_document.py`. Nu
acopereau si adaugarea intr-un act care exista deja, iar acelea sunt doua drumuri
diferite prin toolkit: `docx-footnotes` construieste, `docx-track-changes` editeaza.

Pe 4 august 2026, la 13:34, un act cu antet, A4 si 51 de paragrafe numerotate a primit
treisprezece paragrafe noi despre corespondenta partilor, scrise cu `add_paragraph`. Au
iesit pe `Normal`, cu alineat de prima linie de 1,27 cm si fara numar, in mijlocul unui
act corect. Autorul a crezut ca regula s-a pierdut din nou.

Ce indreapta
------------
1. Paragraful de corp scris pe `Normal` cu alineat de prima linie trece pe stilul de corp,
   primeste numerotarea automata a documentului si indentarea atarnata masurata pe actele
   proprii, `w:ind w:left="0" w:hanging="720"`.
2. Citatul, adica paragraful cu tot textul in italic, primeste forma de bloc retras,
   `w:left="720" w:firstLine="0"`, si ramane in afara fluxului numerotat.
3. Fiecare titlu primeste rand gol deasupra, daca nu il are.

Ce nu atinge
------------
Textul, evidentierile, notele de subsol, antetul, subsolul, marginile. Numarul fluxului
de numerotare se citeste din document, din primul paragraf de corp existent, deci se
foloseste fluxul actului, nu unul presupus.

Utilizare
---------
    python indreapta_forma.py --input act.docx                  # raporteaza, nu scrie
    python indreapta_forma.py --input act.docx --scrie          # indreapta pe loc
    python indreapta_forma.py --input act.docx --output nou.docx
"""
from __future__ import annotations

import argparse
import copy
import json
import shutil
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("Lipseste python-docx. Instaleaza-l cu: pip install python-docx")

AICI = Path(__file__).resolve().parent
sys.path.insert(0, str(AICI))

STIL_CORP = "Body cu nr de paragraf"
STIL_SIMPLU = "Normal"
INDENT_CORP = (0, 720)
INDENT_CITAT = 720

ORDINE_PPR = (
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr", "widowControl",
    "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs", "suppressAutoHyphens",
    "kinsoku", "wordWrap", "overflowPunct", "topLinePunct", "autoSpaceDE", "autoSpaceDN",
    "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind", "contextualSpacing",
    "mirrorIndents", "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr", "pPrChange",
)


def _pune_in_pPr(ppr, element, nume: str):
    for vechi in ppr.findall(qn("w:" + nume)):
        ppr.remove(vechi)
    rang = ORDINE_PPR.index(nume)
    for copil in ppr:
        eticheta = copil.tag.split("}")[-1]
        if eticheta in ORDINE_PPR and ORDINE_PPR.index(eticheta) > rang:
            copil.addprevious(element)
            return
    ppr.append(element)


def _numerotat(p) -> bool:
    ppr = p._p.pPr
    return ppr is not None and ppr.find(qn("w:numPr")) is not None


def _alineat(p):
    """Valoarea `w:firstLine`, sau None. Zero inseamna fara alineat, deci None."""
    ppr = p._p.pPr
    ind = ppr.find(qn("w:ind")) if ppr is not None else None
    val = ind.get(qn("w:firstLine")) if ind is not None else None
    return None if val in (None, "0") else val


def _e_citat(p) -> bool:
    rulate = [r for r in p.runs if r.text.strip()]
    return bool(rulate) and all(r.italic for r in rulate)


_GHILIMELE_DESCHIDERE = ("„", "«", '"')
_GHILIMELE_INCHIDERE = ("”", "»", '"')


def _arata_ca_citat(text: str) -> bool:
    """Un paragraf care incepe si se termina cu ghilimele, indiferent de stil sau de
    italic. Prinde citatul introdus prin insertie de track-changes, care mosteneste
    stilul paragrafului de ancora si iese numerotat, pe `Body cu nr de paragraf`, fara
    sa treaca vreodata prin `Normal`, drumul pe care il verifica `_e_citat`."""
    return bool(text) and text[0] in _GHILIMELE_DESCHIDERE and text[-1] in _GHILIMELE_INCHIDERE


def _e_titlu(p) -> bool:
    return p.style.name.lower().startswith("heading") and bool(p.text.strip())


def numar_flux(doc):
    """Fluxul de numerotare al corpului, citit din document, nu presupus."""
    for p in doc.paragraphs:
        if p.style.name != STIL_CORP or not _numerotat(p):
            continue
        numpr = p._p.pPr.find(qn("w:numPr"))
        nid = numpr.find(qn("w:numId"))
        if nid is not None:
            return int(nid.get(qn("w:val")))
    return None


def _treci_pe_corp(doc, p, num_id: int):
    ppr = p._p.get_or_add_pPr()
    for vechi in ppr.findall(qn("w:pStyle")):
        ppr.remove(vechi)
    st = OxmlElement("w:pStyle")
    st.set(qn("w:val"), doc.styles[STIL_CORP].style_id)
    _pune_in_pPr(ppr, st, "pStyle")

    for vechi in ppr.findall(qn("w:numPr")):
        ppr.remove(vechi)
    npr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    npr.append(ilvl)
    npr.append(nid)
    _pune_in_pPr(ppr, npr, "numPr")

    stanga, atarnare = INDENT_CORP
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(stanga))
    ind.set(qn("w:hanging"), str(atarnare))
    _pune_in_pPr(ppr, ind, "ind")


def _treci_pe_citat(p):
    ppr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(INDENT_CITAT))
    ind.set(qn("w:firstLine"), "0")
    _pune_in_pPr(ppr, ind, "ind")


def _de_pe_corp_pe_citat(doc, p):
    """Un paragraf deja numerotat pe stilul de corp, al carui text e chiar un citat:
    scoate numerotarea, trece pe stilul simplu, aplica indentarea de citat retras si
    italicul pe fiecare rand cu text, ca sa iasa in forma pe care o poarta un citat
    scris de la inceput pe `Normal`."""
    ppr = p._p.get_or_add_pPr()
    for vechi in ppr.findall(qn("w:numPr")):
        ppr.remove(vechi)
    for vechi in ppr.findall(qn("w:pStyle")):
        ppr.remove(vechi)
    st = OxmlElement("w:pStyle")
    st.set(qn("w:val"), doc.styles[STIL_SIMPLU].style_id)
    _pune_in_pPr(ppr, st, "pStyle")
    _treci_pe_citat(p)
    for r in p.runs:
        if r.text.strip():
            r.italic = True


def _rand_gol_inainte(doc, p):
    """Insereaza un paragraf gol inaintea titlului, pe stilul simplu, aliniat la margine."""
    gol = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    st = OxmlElement("w:pStyle")
    st.set(qn("w:val"), doc.styles[STIL_SIMPLU].style_id)
    ppr.append(st)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ppr.append(ind)
    gol.append(ppr)
    p._p.addprevious(gol)
    return gol


def indreapta(cale: Path, scrie: bool) -> dict:
    doc = Document(str(cale))
    num_id = numar_flux(doc)

    trecute_corp, trecute_citat, randuri_goale = [], [], []
    if num_id is not None:
        for p in doc.paragraphs:
            t = p.text.strip()
            if not t or _e_titlu(p):
                continue
            if _numerotat(p):
                # Numerotat, dar textul e chiar un citat: scapat prin insertie de
                # track-changes, care mosteneste stilul paragrafului de ancora in loc
                # sa treaca prin `Normal`, singurul drum verificat mai jos.
                if p.style.name == STIL_CORP and _arata_ca_citat(t):
                    trecute_citat.append(t[:60])
                    if scrie:
                        _de_pe_corp_pe_citat(doc, p)
                continue
            if p.style.name != STIL_SIMPLU:
                continue
            if _e_citat(p):
                if _alineat(p) is not None:
                    trecute_citat.append(t[:60])
                    if scrie:
                        _treci_pe_citat(p)
                continue
            if _alineat(p) is not None:
                trecute_corp.append(t[:60])
                if scrie:
                    _treci_pe_corp(doc, p, num_id)

    # Randul gol dinaintea titlului se pune la sfarsit, ca inserarile sa nu strice mersul
    # de mai sus prin lista de paragrafe.
    parag = doc.paragraphs
    for k, p in enumerate(parag):
        if not _e_titlu(p) or k == 0:
            continue
        if not parag[k - 1].text.strip():
            continue
        randuri_goale.append(p.text.strip()[:50])
        if scrie:
            _rand_gol_inainte(doc, p)

    raport = {
        "fisier": str(cale),
        "flux_numerotare": num_id,
        "paragrafe_trecute_pe_corp": trecute_corp,
        "citate_retrase": trecute_citat,
        "titluri_care_au_primit_rand_gol": randuri_goale,
        "scris": bool(scrie),
    }
    if num_id is None:
        raport["avertisment"] = ("documentul nu are niciun paragraf pe stilul de corp, "
                                 "deci nu e un act in forma casei; nu am ce indrepta")
    return raport, doc


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Indreapta forma unui act deja existent")
    ap.add_argument("--input", required=True, type=Path)
    ap.add_argument("--output", type=Path,
                    help="scrie rezultatul alaturi; fara el se scrie pe loc, cu --scrie")
    ap.add_argument("--scrie", action="store_true",
                    help="fara el doar raporteaza, nu atinge fisierul")
    args = ap.parse_args(argv)

    if not args.input.exists():
        sys.exit("Fisierul nu exista: %s" % args.input)

    scrie = bool(args.scrie or args.output)
    raport, doc = indreapta(args.input, scrie)

    if scrie:
        tinta = args.output or args.input
        if tinta == args.input:
            copie = args.input.with_name(
                args.input.stem + ".inainte-de-indreptare" + args.input.suffix)
            shutil.copy2(args.input, copie)
            raport["copie_de_siguranta"] = str(copie)
        doc.save(str(tinta))
        raport["iesire"] = str(tinta)
        try:
            import repara_pachet  # noqa: PLC0415
            repara_pachet.asigura(tinta, tacut=True)
        except Exception as e:  # noqa: BLE001
            raport["spatii_de_nume"] = "neverificate: %s" % e

    print(json.dumps(raport, ensure_ascii=False, indent=1))
    return 0


if __name__ == "__main__":
    sys.exit(main())
