#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
creeaza_document.py

Construieste un .docx gol in stilul de casa complet al cabinetului, gata de scris.

Motivul pentru care exista. Descrierea in proza a stilului prinde fontul si marginile,
dar pierde ce nu se poate spune in cuvinte: logo-ul din antet, campurile PAGE si
NUMPAGES din subsol, culoarea exacta a titlurilor, definitia de numerotare pe niveluri.
Toate acestea se scriu in XML, deci le construim, nu le explicam.

Valorile vin din masuratoarea pe 86 de acte proprii, cele cu corpul Georgia, separate
de documentele primite pe sabloane straine. Vezi tabelul din SKILL.md.

Utilizare
---------
    python creeaza_document.py --output nota.docx
    python creeaza_document.py --output cerere.docx --titlu "CERERE DE CHEMARE IN JUDECATA"
    python creeaza_document.py --output act.docx --fara-antet
    python creeaza_document.py --output act.docx --academic
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    sys.exit("Lipseste python-docx. Instaleaza-l cu: pip install python-docx")

AICI = Path(__file__).resolve().parent
LOGO = AICI.parent / "assets" / "antet-amz.png"

# Stilul de casa, masurat pe 86 de acte proprii.
CASA = {
    "font": "Georgia",
    "corp_pt": 10,
    "interlinie_pt": 14,          # "cel putin", nu exact
    "spatiu_pt": 6,               # inainte si dupa paragraf
    "alineat_cm": 1.27,
    "margini_cm": {"sus": 2.0, "jos": 0.71, "stanga": 3.5, "dreapta": 1.5},
    "antet_de_sus_cm": 1.96,
    "subsol_de_jos_cm": 0.71,
    "logo_latime_cm": 15.96,
    "titlu_pt": 13,
    "titlu_culoare": "244061",    # bleumarin, 58% din titlurile proprii
    "nota_pt": 8,
    "subsol_pt": 7.5,
}

ACADEMIC = {
    "font": "Times New Roman",
    "corp_pt": 12,
    "interlinie_pt": None,        # 1,5 randuri, nu valoare fixa
    "spatiu_pt": 0,
    "alineat_cm": 1.27,
    "margini_cm": {"sus": 2.5, "jos": 2.5, "stanga": 3.0, "dreapta": 2.0},
    "antet_de_sus_cm": 1.25,
    "subsol_de_jos_cm": 1.25,
    "logo_latime_cm": None,
    "titlu_pt": 14,
    "titlu_culoare": None,        # negru
    "nota_pt": 10,
    "subsol_pt": 11,
}


def _camp(paragraf, instructiune: str):
    """Insereaza un camp Word, de exemplu PAGE sau NUMPAGES.

    python-docx nu are API pentru campuri, iar un numar scris ca text s-ar bloca la
    prima pagina adaugata. Construim cele trei elemente cerute de format: begin,
    instrText, end.
    """
    run = paragraf.add_run()
    inceput = OxmlElement("w:fldChar")
    inceput.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instructiune
    sfarsit = OxmlElement("w:fldChar")
    sfarsit.set(qn("w:fldCharType"), "end")
    run._r.append(inceput)
    run._r.append(instr)
    run._r.append(sfarsit)
    return run


def _font_complet(element, nume: str):
    """Declara fontul si pentru scriptul est-european, altfel Word cade pe substitut
    la diacriticele romanesti scrise cu virgula dedesubt."""
    rpr = element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for atribut in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(atribut), nume)


def _stil_normal(doc, ref):
    normal = doc.styles["Normal"]
    normal.font.name = ref["font"]
    normal.font.size = Pt(ref["corp_pt"])
    _font_complet(normal.element, ref["font"])

    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(ref["alineat_cm"])
    pf.space_before = Pt(ref["spatiu_pt"])
    pf.space_after = Pt(ref["spatiu_pt"])
    if ref["interlinie_pt"]:
        pf.line_spacing = Pt(ref["interlinie_pt"])
        # Ordinea conteaza: line_spacing cu o lungime scrie lineRule="exact", care taie
        # randul la un indice sau la un caracter mai inalt. Il corectam in "atLeast".
        pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def _stiluri_titlu(doc, ref):
    """Titlurile proprii sunt bold, majuscule, centrate, bleumarin."""
    for nivel, marime in ((1, ref["titlu_pt"]), (2, ref["titlu_pt"] - 1), (3, ref["titlu_pt"] - 2)):
        try:
            stil = doc.styles[f"Heading {nivel}"]
        except KeyError:
            continue
        stil.font.name = ref["font"]
        stil.font.size = Pt(marime)
        stil.font.bold = True
        _font_complet(stil.element, ref["font"])
        if ref["titlu_culoare"]:
            stil.font.color.rgb = RGBColor.from_string(ref["titlu_culoare"])
        pf = stil.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if nivel == 1 else WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.keep_with_next = True
        # Majuscule afisate prin stil, nu scrise in text: cine copiaza titlul primeste
        # forma originala, iar un redline nu marcheaza diferenta de capitalizare.
        rpr = stil.element.get_or_add_rPr()
        caps = OxmlElement("w:caps")
        caps.set(qn("w:val"), "1")
        rpr.append(caps)


def _pagina(doc, ref):
    for s in doc.sections:
        s.top_margin = Cm(ref["margini_cm"]["sus"])
        s.bottom_margin = Cm(ref["margini_cm"]["jos"])
        s.left_margin = Cm(ref["margini_cm"]["stanga"])
        s.right_margin = Cm(ref["margini_cm"]["dreapta"])
        s.header_distance = Cm(ref["antet_de_sus_cm"])
        s.footer_distance = Cm(ref["subsol_de_jos_cm"])


def _antet(doc, ref, cu_logo: bool):
    if not cu_logo or not ref["logo_latime_cm"]:
        return False
    if not LOGO.exists():
        print(f"  atentie: lipseste {LOGO}, antetul ramane gol", file=sys.stderr)
        return False
    for s in doc.sections:
        p = s.header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(LOGO), width=Cm(ref["logo_latime_cm"]))
    return True


def _subsol(doc, ref):
    """Pagina N din M, aliniat dreapta, cu campuri reale."""
    for s in doc.sections:
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run("Pagina ")
        _camp(p, r" PAGE \* MERGEFORMAT ")
        r2 = p.add_run(" din ")
        _camp(p, r" NUMPAGES \* MERGEFORMAT ")
        for run in p.runs:
            run.font.name = ref["font"]
            run.font.size = Pt(ref["subsol_pt"])
            _font_complet(run.element, ref["font"])


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Document nou in stilul de casa")
    ap.add_argument("--output", required=True, type=Path)
    ap.add_argument("--titlu", help="titlul actului, scris ca Heading 1")
    ap.add_argument("--fara-antet", action="store_true", help="document fara logo")
    ap.add_argument("--academic", action="store_true",
                    help="format Scoala Doctorala UB Drept in loc de stilul de casa")
    ap.add_argument("--autor", default="Adrian Zamfir")
    args = ap.parse_args(argv)

    ref = ACADEMIC if args.academic else CASA
    doc = Document()

    _pagina(doc, ref)
    _stil_normal(doc, ref)
    _stiluri_titlu(doc, ref)
    cu_logo = _antet(doc, ref, not args.fara_antet)
    _subsol(doc, ref)

    if args.titlu:
        doc.add_heading(args.titlu, level=1)

    core = doc.core_properties
    core.author = args.autor
    core.last_modified_by = args.autor

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))

    print(f"{args.output}")
    print(f"  profil: {'academic UB Drept' if args.academic else 'stil de casa'}")
    print(f"  antet cu logo: {'da' if cu_logo else 'nu'}")
    print(f"  subsol: Pagina N din M, {ref['font']} {ref['subsol_pt']} pt")
    print(f"  titluri: {ref['font']} {ref['titlu_pt']} pt bold caps"
          f"{', #' + ref['titlu_culoare'] if ref['titlu_culoare'] else ''}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
