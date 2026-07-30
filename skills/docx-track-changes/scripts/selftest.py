# -*- coding: utf-8 -*-
"""Self-test pentru docx_track_changes.py.

Construieste documente de proba, aplica revizii si verifica invariantele:
  1. Reject All reproduce exact documentul original.
  2. Accept All produce exact varianta ceruta.
  3. Formatarea run-urilor (bold, italic) supravietuieste editarii.
  4. Referintele de nota de subsol raman la locul lor.
  5. Editarile care nu se potrivesc opresc executia, fara fisier de iesire.
  6. Marcajul acopera cuvantul schimbat, nu fraza din jurul lui.

Rulare:  python selftest.py
"""

from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import tempfile
import zipfile
from copy import deepcopy

from lxml import etree

HERE = os.path.dirname(os.path.abspath(__file__))
TOOL = os.path.join(HERE, "docx_track_changes.py")

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XMLSP = "{http://www.w3.org/XML/1998/namespace}space"


def w(t):
    return "{%s}%s" % (W, t)


FAILURES = []


def check(cond, label):
    print(("  OK   " if cond else "  FAIL ") + label)
    if not cond:
        FAILURES.append(label)


def run_tool(*args):
    proc = subprocess.run([sys.executable, TOOL] + list(args),
                          capture_output=True, text=True, encoding="utf-8")
    return proc.returncode, proc.stdout, proc.stderr


# --------------------------------------------------------------------------

def build_sample(path):
    from docx import Document
    d = Document()
    d.add_heading("Cerere de sesizare a Curtii Constitutionale", level=1)
    p = d.add_paragraph()
    p.add_run("Subsemnatul, ")
    p.add_run("Adrian Zamfir").bold = True
    p.add_run(", avocat, formulez prezenta ")
    p.add_run("exceptie de neconstitutionalitate").italic = True
    p.add_run(" impotriva art. 12 din Legea nr. 58/2023.")
    d.add_paragraph("In motivare, aratam ca textul criticat incalca exigentele de "
                    "previzibilitate impuse de art. 1 alin. (5) din Constitutie.")
    d.add_paragraph("Paragraf care va fi sters integral.")
    d.add_paragraph("Paragraf martor, ramane neatins.")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "Temei juridic"
    t.cell(0, 1).text = "art. 29 din Legea nr. 47/1992"
    d.save(path)


def build_sample_footnote(path, tmp):
    from docx import Document
    base = os.path.join(tmp, "_base.docx")
    d = Document()
    d.add_paragraph("Textul criticat incalca art. 1 alin. (5) din Constitutie")
    d.save(base)

    footnotes = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:footnotes xmlns:w="%s">'
        '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
        '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r>'
        '<w:continuationSeparator/></w:r></w:p></w:footnote>'
        '<w:footnote w:id="1"><w:p><w:r><w:footnoteRef/></w:r><w:r>'
        '<w:t xml:space="preserve"> C.C.R., Decizia nr. 70/2023, par. 65.</w:t>'
        '</w:r></w:p></w:footnote></w:footnotes>' % W
    ).encode("utf-8")

    with zipfile.ZipFile(base) as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}

    doc = etree.fromstring(data["word/document.xml"])
    first = list(doc.find(w("body")).iter(w("p")))[0]
    ref = etree.SubElement(first, w("r"))
    rpr = etree.SubElement(ref, w("rPr"))
    etree.SubElement(rpr, w("rStyle")).set(w("val"), "FootnoteReference")
    etree.SubElement(ref, w("footnoteReference")).set(w("id"), "1")
    tail = etree.SubElement(first, w("r"))
    tt = etree.SubElement(tail, w("t"))
    tt.set(XMLSP, "preserve")
    tt.text = ", sub aspectul previzibilitatii."
    data["word/document.xml"] = etree.tostring(doc, xml_declaration=True,
                                               encoding="UTF-8", standalone=True)
    data["word/footnotes.xml"] = footnotes

    ct = etree.fromstring(data["[Content_Types].xml"])
    ov = etree.SubElement(ct, "{%s}Override" % ct.nsmap[None])
    ov.set("PartName", "/word/footnotes.xml")
    ov.set("ContentType", "application/vnd.openxmlformats-officedocument."
                          "wordprocessingml.footnotes+xml")
    data["[Content_Types].xml"] = etree.tostring(ct, xml_declaration=True,
                                                 encoding="UTF-8", standalone=True)

    rels = etree.fromstring(data["word/_rels/document.xml.rels"])
    rel = etree.SubElement(rels, "{%s}Relationship" % rels.nsmap[None])
    rel.set("Id", "rIdFootnotesTest")
    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/"
                    "relationships/footnotes")
    rel.set("Target", "footnotes.xml")
    data["word/_rels/document.xml.rels"] = etree.tostring(rels, xml_declaration=True,
                                                          encoding="UTF-8", standalone=True)

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, data[n])
        zout.writestr("word/footnotes.xml", data["word/footnotes.xml"])
    os.remove(base)


# --------------------------------------------------------------------------

def doc_root(path):
    with zipfile.ZipFile(path) as z:
        return etree.fromstring(z.read("word/document.xml"))


def plain_texts(root):
    body = root.find(w("body"))
    return ["".join(t.text or "" for t in p.iter(w("t"))) for p in body.iter(w("p"))]


def resolve(root, mode):
    """mode 'accept' sau 'reject'; intoarce textele de paragraf rezultate."""
    r = deepcopy(root)
    drop = w("ins") if mode == "reject" else w("del")
    keep = w("del") if mode == "reject" else w("ins")
    for el in list(r.iter(drop)):
        parent = el.getparent()
        if parent is not None and parent.tag != w("rPr"):
            parent.remove(el)
    for el in list(r.iter(keep)):
        parent = el.getparent()
        if parent is None or parent.tag == w("rPr"):
            continue
        idx = list(parent).index(el)
        for c in reversed(list(el)):
            parent.insert(idx, c)
        parent.remove(el)
    for dt in list(r.iter(w("delText"))):
        dt.tag = w("t")
    out = []
    for p in r.find(w("body")).iter(w("p")):
        ppr = p.find(w("pPr"))
        marked = ppr is not None and ppr.find(w("rPr")) is not None \
            and ppr.find(w("rPr")).find(drop) is not None
        txt = "".join(t.text or "" for t in p.iter(w("t")))
        if marked and not txt.strip():
            continue
        out.append(txt)
    return out


# --------------------------------------------------------------------------

def main():
    tmp = tempfile.mkdtemp(prefix="dtc_selftest_")
    try:
        orig = os.path.join(tmp, "original.docx")
        out = os.path.join(tmp, "revizuit.docx")
        build_sample(orig)

        edits = {
            "author": "Test",
            "edits": [
                {"op": "replace", "find": "art. 12 din Legea nr. 58/2023",
                 "replace": "art. 12 alin. (2) din Legea nr. 58/2023"},
                {"op": "replace", "find": "exigentele de previzibilitate",
                 "replace": "exigentele de claritate si previzibilitate"},
                {"op": "delete_paragraph", "id": "doc:3"},
                {"op": "insert_after", "id": "doc:2",
                 "text": "Invocam Decizia nr. 70/2023 a Curtii Constitutionale."},
                {"op": "set_paragraph", "id": "doc:6",
                 "text": "art. 29 alin. (1) din Legea nr. 47/1992"},
            ],
        }
        ep = os.path.join(tmp, "edits.json")
        with open(ep, "w", encoding="utf-8") as fh:
            json.dump(edits, fh, ensure_ascii=False)

        before = plain_texts(doc_root(orig))
        code, so, se = run_tool("apply", "--input", orig, "--output", out, "--edits", ep)
        print("\n[1] editari punctuale prin edits.json")
        check(code == 0, "apply se termina cu succes")
        if code != 0:
            print(se)
            return 1

        root = doc_root(out)
        check(resolve(root, "reject") == before, "Reject All reproduce exact originalul")
        expected = [
            "Cerere de sesizare a Curtii Constitutionale",
            "Subsemnatul, Adrian Zamfir, avocat, formulez prezenta exceptie de "
            "neconstitutionalitate impotriva art. 12 alin. (2) din Legea nr. 58/2023.",
            "In motivare, aratam ca textul criticat incalca exigentele de claritate si "
            "previzibilitate impuse de art. 1 alin. (5) din Constitutie.",
            "Invocam Decizia nr. 70/2023 a Curtii Constitutionale.",
            "Paragraf martor, ramane neatins.",
            "Temei juridic",
            "art. 29 alin. (1) din Legea nr. 47/1992",
        ]
        got = [t for t in resolve(root, "accept") if t.strip()]
        check(got == expected, "Accept All produce exact varianta ceruta")
        if got != expected:
            for a, b in zip(got, expected):
                if a != b:
                    print("     got: %r\n     exp: %r" % (a, b))

        p1 = list(root.find(w("body")).iter(w("p")))[1]
        bold, ital = [], []
        for r in p1.iter(w("r")):
            rpr = r.find(w("rPr"))
            txt = "".join((t.text or "") for t in r.iter(w("t"), w("delText")))
            if rpr is None or not txt:
                continue
            if rpr.find(w("b")) is not None:
                bold.append(txt)
            if rpr.find(w("i")) is not None:
                ital.append(txt)
        check("Adrian Zamfir" in "".join(bold), "bold pastrat in paragraful editat")
        check("exceptie de neconstitutionalitate" in "".join(ital),
              "italic pastrat in paragraful editat")

        marks = [p for p in root.find(w("body")).iter(w("p"))
                 if p.find(w("pPr")) is not None
                 and p.find(w("pPr")).find(w("rPr")) is not None]
        ok_order = True
        for p in marks:
            ppr = p.find(w("pPr"))
            kids = [c.tag for c in ppr]
            if w("sectPr") in kids and kids.index(w("rPr")) > kids.index(w("sectPr")):
                ok_order = False
        check(ok_order, "ordinea elementelor in w:pPr respecta schema OOXML")

        print("\n[2] varianta revizuita ca fisier text")
        md = os.path.join(tmp, "revizuit.md")
        with open(md, "w", encoding="utf-8") as fh:
            fh.write("\n".join(expected))
        out2 = os.path.join(tmp, "rev_md.docx")
        code, so, se = run_tool("apply", "--input", orig, "--output", out2, "--revised", md)
        check(code == 0, "apply cu --revised se termina cu succes")
        if code == 0:
            r2 = doc_root(out2)
            check(resolve(r2, "reject") == before, "Reject All reproduce originalul")
            check([t for t in resolve(r2, "accept") if t.strip()] == expected,
                  "Accept All produce varianta din fisierul text")

        print("\n[3] note de subsol")
        fn_orig = os.path.join(tmp, "original_fn.docx")
        fn_out = os.path.join(tmp, "revizuit_fn.docx")
        build_sample_footnote(fn_orig, tmp)
        fn_edits = [
            {"op": "replace", "find": "art. 1 alin. (5) din Constitutie",
             "replace": "art. 1 alin. (5) teza a doua din Constitutie"},
            {"op": "replace", "find": "sub aspectul previzibilitatii",
             "replace": "sub aspectul calitatii legii"},
            {"op": "replace", "part": "footnotes", "find": "par. 65.",
             "replace": "par. 65-67."},
        ]
        fep = os.path.join(tmp, "edits_fn.json")
        with open(fep, "w", encoding="utf-8") as fh:
            json.dump(fn_edits, fh, ensure_ascii=False)
        code, so, se = run_tool("apply", "--input", fn_orig, "--output", fn_out, "--edits", fep)
        check(code == 0, "apply pe document cu note se termina cu succes")
        if code == 0:
            r3 = doc_root(fn_out)
            refs = list(r3.iter(w("footnoteReference")))
            check(len(refs) == 1 and refs[0].get(w("id")) == "1",
                  "referinta de nota de subsol a supravietuit editarii")
            with zipfile.ZipFile(fn_out) as z:
                fnr = etree.fromstring(z.read("word/footnotes.xml"))
            check(len(list(fnr.iter(w("footnote")))) == 3, "toate notele raman in footnotes.xml")
            check(fnr.find(".//" + w("ins")) is not None,
                  "revizia din nota de subsol a fost scrisa")

        print("\n[4] editari care nu se potrivesc")
        bad = os.path.join(tmp, "bad.json")
        with open(bad, "w", encoding="utf-8") as fh:
            json.dump([{"op": "replace", "find": "sir inexistent xyzzy",
                        "replace": "x"}], fh)
        nope = os.path.join(tmp, "nope.docx")
        code, so, se = run_tool("apply", "--input", orig, "--output", nope, "--edits", bad)
        check(code != 0, "apply esueaza cand sirul cautat lipseste")
        check(not os.path.exists(nope), "nu scrie fisier de iesire la esec")
        check("nu exista in document" in se, "mesajul de eroare spune ce nu s-a potrivit")

        bad2 = os.path.join(tmp, "bad2.json")
        with open(bad2, "w", encoding="utf-8") as fh:
            json.dump([{"op": "set_paragraph", "id": "doc:999", "text": "x"}], fh)
        code, so, se = run_tool("apply", "--input", orig, "--output", nope, "--edits", bad2)
        check(code != 0 and "in afara intervalului" in se,
              "apply esueaza la un id de paragraf inexistent")

        print("\n[5] diacritice scrise altfel, aceeasi litera")
        # Sedila (ş U+015F, ţ U+0163) si virgula dedesubt (ș U+0219, ț U+021B) arata
        # identic. Un citat preluat dintr-o baza de legislatie aduce sedila, restul
        # actului are virgula, iar o conversie de codificare inecase documentul in
        # revizii pe cuvinte care se citesc la fel. Aici nu trebuie sa apara niciuna.
        import docx as _docx
        dia_orig = os.path.join(tmp, "dia_orig.docx")
        d = _docx.Document()
        d.add_paragraph("Instanţa a reţinut că deţinerea şi portul armei sunt supuse autorizării.")
        d.add_paragraph("Al doilea paragraf ramane neatins.")
        d.save(dia_orig)

        dia_rev = os.path.join(tmp, "dia_rev.txt")
        with open(dia_rev, "w", encoding="utf-8") as fh:
            # aceleasi cuvinte, scrise cu virgula dedesubt, plus o modificare reala
            fh.write("Instanța a reținut că deținerea și portul armei sunt supuse autorizării.\n")
            fh.write("Al doilea paragraf a fost modificat.\n")

        dia_out = os.path.join(tmp, "dia_out.docx")
        code, so, se = run_tool("apply", "--input", dia_orig, "--output", dia_out,
                                "--revised", dia_rev, "--author", "AMZ Law Office")
        check(code == 0, "apply se termina cu succes pe diacritice amestecate")
        if code == 0:
            dr = doc_root(dia_out)
            paras = list(dr.iter(w("p")))
            p0_ins = list(paras[0].iter(w("ins")))
            p0_del = list(paras[0].iter(w("del")))
            check(not p0_ins and not p0_del,
                  "sedila fata de virgula dedesubt nu produce nicio revizie")
            txt0 = "".join(t.text or "" for t in paras[0].iter(w("t")))
            check("Instanţa" in txt0 and "reţinut" in txt0,
                  "paragraful nemodificat isi pastreaza grafia din original")
            check(list(paras[1].iter(w("ins"))), "modificarea reala se marcheaza in continuare")

        print("\n[6] forma afisata de Word, nu textul stocat")
        # Word afiseaza numarul de lista din w:numPr si majusculele din w:caps, fara
        # ca ele sa existe in text. O varianta revizuita venita dintr-o retiparire le
        # aduce inapoi ca text. Cititorul vede acelasi lucru, deci nu se cuvine nicio
        # revizie, iar titlul trebuie sa-si pastreze stilul, nu sa fie sters si rescris.
        import docx as _dx
        from docx.enum.style import WD_STYLE_TYPE as _ST

        vz_orig = os.path.join(tmp, "vizual_orig.docx")
        doc = _dx.Document()
        stil = doc.styles.add_style("TitluCaps", _ST.PARAGRAPH)
        stil.font.all_caps = True
        h = doc.add_paragraph("Incalcarea articolului 1 din Constitutie", style="TitluCaps")
        # numerotare automata pe titlu
        ppr = h._p.get_or_add_pPr()
        numpr = ppr.makeelement(w("numPr"), {})
        ilvl = ppr.makeelement(w("ilvl"), {w("val"): "0"})
        numid = ppr.makeelement(w("numId"), {w("val"): "1"})
        numpr.append(ilvl)
        numpr.append(numid)
        ppr.append(numpr)
        doc.add_paragraph("Un paragraf oarecare, neatins.")
        doc.save(vz_orig)

        vz_rev = os.path.join(tmp, "vizual_rev.txt")
        with open(vz_rev, "w", encoding="utf-8") as fh:
            # exact ce iese dintr-o retiparire: numarul scris ca text, titlul cu
            # majuscule, plus doua cuvinte chiar adaugate
            fh.write("IV. INCALCAREA ARTICOLULUI 1 DIN CONSTITUTIE SI DIN CEDO\n")
            fh.write("Un paragraf oarecare, neatins.\n")

        vz_out = os.path.join(tmp, "vizual_out.docx")
        code, so, se = run_tool("apply", "--input", vz_orig, "--output", vz_out,
                                "--revised", vz_rev, "--author", "AMZ Law Office")
        check(code == 0, "apply se termina cu succes pe titlu numerotat si cu majuscule")
        if code == 0:
            vr = doc_root(vz_out)
            titlu = list(vr.iter(w("p")))[0]
            stil_ramas = titlu.find(w("pPr") + "/" + w("pStyle"))
            check(stil_ramas is not None and stil_ramas.get(w("val")) == "TitluCaps",
                  "titlul isi pastreaza stilul, nu e sters si rescris ca paragraf nou")
            check(titlu.find(w("pPr") + "/" + w("numPr")) is not None,
                  "titlul isi pastreaza numerotarea automata")
            sters = "".join(t.text or "" for t in titlu.iter(w("delText")))
            check(not sters.strip(), "nimic nu se sterge din titlu")
            adaugat = "".join(t.text or ""
                              for e in titlu.iter(w("ins")) for t in e.iter(w("t")))
            check("IV" not in adaugat,
                  "numarul de lista scris ca text nu se insereaza peste cel automat")
            check("CEDO" in adaugat.upper(),
                  "cuvintele chiar adaugate se marcheaza in continuare")
            intreg = "".join(t.text or "" for t in titlu.iter(w("t")))
            check("Incalcarea articolului" in intreg,
                  "textul titlului ramane scris ca in original")

        print("\n[7] verify")
        code, so, se = run_tool("verify", "--input", out, "--detail")
        check(code == 0, "verify se termina cu succes")
        if code == 0:
            rep = json.loads(so)
            check(rep["insertions"] >= 3 and rep["deletions"] >= 1,
                  "verify raporteaza inserarile si stergerile")

        print("\n[8] latimea marcajului, cuvantul schimbat, nu fraza din jur")
        # Cand dintr-o paranteza e gresit un singur cuvant, avocatul trebuie sa
        # vada marcat acel cuvant. O paranteza taiata intreaga si rescrisa
        # intreaga il pune sa citeasca de doua ori acelasi text ca sa gaseasca
        # litera schimbata. Editarea se scrie ca fraza, marcajul iese pe cuvant.
        import docx as _dx8
        lat_orig = os.path.join(tmp, "latime_orig.docx")
        d8 = _dx8.Document()
        d8.add_paragraph("Deşi Raspunsul nu a fost motivat in drept (fiind nelegal "
                         "pentru aceste aspect), instanta va constata nelegalitatea.")
        d8.save(lat_orig)

        e8 = os.path.join(tmp, "edits_latime.json")
        with open(e8, "w", encoding="utf-8") as fh:
            json.dump([{"op": "replace",
                        "find": "(fiind nelegal pentru aceste aspect)",
                        "replace": "(fiind nelegal pentru acest aspect)"}],
                      fh, ensure_ascii=False)
        lat_out = os.path.join(tmp, "latime_out.docx")
        code, so, se = run_tool("apply", "--input", lat_orig, "--output", lat_out,
                                "--edits", e8)
        check(code == 0, "apply se termina cu succes pe o editare scrisa ca fraza")
        if code == 0:
            r8 = doc_root(lat_out)
            sterse = [s for s in ("".join(t.text or "" for t in e.iter(w("delText")))
                                  for e in r8.iter(w("del"))) if s]
            adaugate = [s for s in ("".join(t.text or "" for t in e.iter(w("t")))
                                    for e in r8.iter(w("ins"))) if s]
            check(sterse == ["aceste"], "se sterge doar cuvantul gresit, nu toata paranteza")
            check(adaugate == ["acest"], "se insereaza doar cuvantul corect")

            # acelasi continut, marcat lat, asa cum iese dintr-o editare pe XML
            # despachetat care inlocuieste un <w:r> intreg
            lat_manual = os.path.join(tmp, "latime_manual.docx")
            with zipfile.ZipFile(lat_orig) as zin:
                names8 = zin.namelist()
                data8 = {n: zin.read(n) for n in names8}
            doc8 = etree.fromstring(data8["word/document.xml"])
            p8 = list(doc8.iter(w("p")))[0]
            for r in list(p8):
                if r.tag == w("r"):
                    p8.remove(r)
            for tag, tt, text in (("del", "delText", "(fiind nelegal pentru aceste aspect)"),
                                  ("ins", "t", "(fiind nelegal pentru acest aspect)")):
                holder = etree.SubElement(p8, w(tag))
                holder.set(w("id"), "5001" if tag == "del" else "5002")
                holder.set(w("author"), "Test")
                holder.set(w("date"), "2026-07-29T00:00:00Z")
                run = etree.SubElement(holder, w("r"))
                node = etree.SubElement(run, w(tt))
                node.set(XMLSP, "preserve")
                node.text = text
            data8["word/document.xml"] = etree.tostring(doc8, xml_declaration=True,
                                                        encoding="UTF-8", standalone=True)
            with zipfile.ZipFile(lat_manual, "w", zipfile.ZIP_DEFLATED) as zout:
                for n in names8:
                    zout.writestr(n, data8[n])

            code, so, se = run_tool("verify", "--input", lat_manual, "--strict")
            check(code != 0, "verify --strict cade pe un marcaj mai lat decat modificarea")
            rep8 = json.loads(so).get("redline_prea_larg") or []
            check(bool(rep8) and rep8[0]["ar_fi_fost_destul"]["sters"] == "aceste",
                  "raportul spune ce cuvant ar fi fost destul sa fie marcat")

            code, so, se = run_tool("verify", "--input", lat_out, "--strict")
            check(code == 0 and json.loads(so)["redline_prea_larg"] == [],
                  "redlineul ingustat de tool trece verificarea stricta")

    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    print("\n" + ("=" * 60))
    if FAILURES:
        print("ESEC: %d verificari picate" % len(FAILURES))
        for f in FAILURES:
            print("  - " + f)
        return 1
    print("Toate verificarile au trecut.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
