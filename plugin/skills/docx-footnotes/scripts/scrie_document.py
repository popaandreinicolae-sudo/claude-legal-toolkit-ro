#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
scrie_document.py

Scrie un act complet in forma casei, pornind de la sablon si de la un continut dat.

De ce exista. Sablonul aduce definitiile de stil, dar cine scrie continutul poate ignora
tot. Documentele generate ieseau cu marginile corecte si structura altcuiva, pe `Normal`
in loc de stilul de corp, cu numerotare tastata sau lipsa, aproape fara bold.

Aici formatarea nu mai e la latitudinea celui care scrie. Textul si rolul fiecarui bloc
vin din afara, iar stilul, numerotarea si bold-ul se aplica de aici, dupa masuratoarea pe
cele 86 de acte proprii.

    python scrie_document.py --continut act.json --output cerere.docx

Formatul continutului
---------------------
{
  "titlu": "cerere de chemare in judecata",
  "instanta": ["CATRE: TRIBUNALUL ARGES", "SECTIA DE CONTENCIOS ADMINISTRATIV SI FISCAL"],
  "reclamant": ["Subsemnata, X, cu domiciliul in ...",
                "Avand domiciliul procesual ales, in temeiul art. 158 C. pr. civ., ..."],
  "parti": ["In contradictoriu cu",
            {"text": "INSPECTORATUL DE POLITIE AL JUDETULUI ARGES", "bold": true},
            "in temeiul Legii nr. 554/2004 a contenciosului administrativ, formulez"],
  "obiect": ["impotriva avizului negativ nr. 384746 din 08.12.2025"],
  "corp": [
    "Paragraf obisnuit, primeste automat numarul [1].",
    {"text": "Termenul de arma interzisa nu este definit.", "bold": ["arma interzisa"]},
    {"tip": "marcator", "text": "element de enumerare, primeste liniuta"},
    {"tip": "titlu", "text": "I. IN FAPT"}
  ],
  "final": ["Cu deosebita consideratie,", "Av. Adrian Zamfir"]
}

Orice bloc lipseste, se sare. `bold` primeste fie true pentru tot paragraful, fie o lista
de fragmente de evidentiat in text.
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("Lipseste python-docx. Instaleaza-l cu: pip install python-docx")

AICI = Path(__file__).resolve().parent
SABLON = AICI.parent / "assets" / "sablon-casa.docx"

# Masurat pe actele proprii. Corpul se numeroteaza [1], [2], [3], cu paranteze drepte,
# prin fluxul 41; enumerarile din interiorul corpului folosesc liniuta, fluxul 37.
# numId=0 anuleaza numerotarea mostenita din stil, folosit la blocul partilor.
NUM_CORP = 41
NUM_MARCATOR = 37
NUM_FARA = 0

# python-docx cere numele stilului, nu identificatorul. Cautarea dupa id e depreciata si
# scoate un avertisment la fiecare paragraf.
STIL_CORP = "Body cu nr de paragraf"
STIL_PARTI = "Parties"
STIL_TITLU = "heading 1"
STIL_SIMPLU = "Normal"


def _numerotare(paragraf, num_id: int, nivel: int = 0):
    """Ataseaza fluxul de numerotare direct pe paragraf.

    Numerele nu se scriu niciodata in text. Unul tastat se strica la prima insertie si
    apare ca modificare intr-un redline, desi cititorul vede acelasi lucru.
    """
    ppr = paragraf._p.get_or_add_pPr()
    for vechi in ppr.findall(qn("w:numPr")):
        ppr.remove(vechi)
    npr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(nivel))
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    npr.append(ilvl)
    npr.append(nid)
    ppr.append(npr)


def _scrie(paragraf, text: str, bold):
    """Aseaza textul, cu evidentierile cerute.

    `bold` poate fi True pentru tot paragraful sau o lista de fragmente. Fragmentele se
    cauta in ordine si textul se sparge in run-uri, ca evidentierea sa cada exact pe ele.
    """
    if bold is True:
        paragraf.add_run(text).bold = True
        return
    if not bold:
        paragraf.add_run(text)
        return

    ramas = text
    for fragment in bold:
        if not fragment or fragment not in ramas:
            continue
        inainte, _, ramas = ramas.partition(fragment)
        if inainte:
            paragraf.add_run(inainte)
        paragraf.add_run(fragment).bold = True
    if ramas:
        paragraf.add_run(ramas)


def _bloc(doc, elemente, stil, *, bold_implicit=False, num_id=None):
    for e in elemente or []:
        item = {"text": e} if isinstance(e, str) else dict(e)
        text = item.get("text", "")
        if not text:
            continue

        tip = item.get("tip", "")
        stil_efectiv = stil
        numerotare = num_id
        if tip == "marcator":
            numerotare = NUM_MARCATOR
        elif tip == "titlu":
            stil_efectiv = STIL_TITLU
            numerotare = None

        try:
            p = doc.add_paragraph(style=stil_efectiv)
        except KeyError:
            p = doc.add_paragraph()

        bold = item.get("bold", bold_implicit)
        _scrie(p, text, bold)
        if numerotare is not None:
            _numerotare(p, numerotare)
        yield p


def construieste(spec: dict, autor: str):
    if not SABLON.exists():
        sys.exit(f"Lipseste sablonul: {SABLON}")
    doc = Document(str(SABLON))

    # Sablonul pastreaza un paragraf gol, ca documentul sa fie valid. Il scoatem dupa ce
    # avem primul continut, altfel actul incepe cu un rand liber.
    gol = doc.paragraphs[0] if doc.paragraphs else None

    numarate = 0
    # Antetul de adresare, identificarea si obiectul stau pe stilul simplu, cu bold.
    list(_bloc(doc, spec.get("instanta"), STIL_SIMPLU, bold_implicit=True))
    list(_bloc(doc, spec.get("reclamant"), STIL_SIMPLU, bold_implicit=True))
    # Blocul partilor isi anuleaza numerotarea mostenita din stil.
    list(_bloc(doc, spec.get("parti"), STIL_PARTI, num_id=NUM_FARA))

    if spec.get("titlu"):
        try:
            p = doc.add_paragraph(style=STIL_TITLU)
        except KeyError:
            p = doc.add_paragraph()
        p.add_run(spec["titlu"]).bold = True

    list(_bloc(doc, spec.get("obiect"), STIL_SIMPLU, bold_implicit=True))

    for p in _bloc(doc, spec.get("corp"), STIL_CORP, num_id=NUM_CORP):
        numarate += 1

    list(_bloc(doc, spec.get("final"), STIL_SIMPLU))

    if gol is not None and len(doc.paragraphs) > 1:
        gol._element.getparent().remove(gol._element)

    core = doc.core_properties
    core.author = autor
    core.last_modified_by = autor
    return doc, numarate


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Act complet in forma casei")
    ap.add_argument("--continut", required=True, type=Path, help="fisier JSON cu continutul")
    ap.add_argument("--output", required=True, type=Path)
    ap.add_argument("--autor", default="Adrian Zamfir")
    args = ap.parse_args(argv)

    spec = json.loads(args.continut.read_text(encoding="utf-8"))
    doc, numarate = construieste(spec, args.autor)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))

    cu_text = sum(1 for p in doc.paragraphs if p.text.strip())
    cu_bold = sum(1 for p in doc.paragraphs if p.text.strip() and any(r.bold for r in p.runs))
    print(f"{args.output}")
    print(f"  paragrafe cu text: {cu_text}, din care numerotate automat: {numarate}")
    print(f"  paragrafe cu bold: {cu_bold} ({100 * cu_bold // max(cu_text, 1)}%), "
          f"in actele proprii 53%")
    print(f"  verifica cu: python ../../docx-livrare-check/scripts/check_livrare.py "
          f"\"{args.output}\"")
    return 0


if __name__ == "__main__":
    sys.exit(main())
