#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
creeaza_document.py

Deschide un document nou pornind de la sablonul de casa, nu de la zero.

De ce asa. Prima varianta reconstruia formatul din parametri masurati, font, marimi,
margini, culori. Suna corect si e gresit: un .docx poarta mult mai mult decat se poate
masura, definitiile complete de stil, numerotarea pe niveluri, tema cu paleta ei,
relatiile dintre sectiune si anteturi, `w:titlePg`. Reconstructia a pierdut exact
`titlePg`, deci logo-ul se repeta pe fiecare pagina in loc sa apara doar pe prima.

Sablonul `assets/sablon-casa.docx` este un act real al cabinetului, golit de continut de
`tools/build_sablon.py`. Pornind de la el, formatul vine intreg, prin constructie.

Utilizare
---------
    python creeaza_document.py --output nota.docx
    python creeaza_document.py --output cerere.docx --titlu "CERERE DE CHEMARE IN JUDECATA"
    python creeaza_document.py --output intern.docx --fara-antet
    python creeaza_document.py --output articol.docx --academic
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError:
    sys.exit("Lipseste python-docx. Instaleaza-l cu: pip install python-docx")

AICI = Path(__file__).resolve().parent
SABLON = AICI.parent / "assets" / "sablon-casa.docx"

sys.path.insert(0, str(AICI))
import repara_pachet  # noqa: E402  spatiile de nume se verifica dupa fiecare salvare

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Formatul academic nu are sablon, fiindca nu exista in corpusul propriu. Ramane singurul
# caz construit din parametri.
ACADEMIC = {
    "font": "Times New Roman",
    "corp_pt": 12,
    "alineat_cm": 1.27,
    "margini_cm": {"sus": 2.5, "jos": 2.5, "stanga": 3.0, "dreapta": 2.0},
    "titlu_pt": 14,
}


def _font_complet(element, nume: str):
    """Declara fontul si pentru scriptul est-european, altfel Word cade pe substitut la
    diacriticele romanesti scrise cu virgula dedesubt."""
    rpr = element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for atribut in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(atribut), nume)


def din_sablon(fara_antet: bool):
    if not SABLON.exists():
        sys.exit(f"Lipseste sablonul: {SABLON}\n"
                 f"Reconstruieste-l cu: python tools/build_sablon.py --sursa <act.docx>")
    doc = Document(str(SABLON))

    if fara_antet:
        # Scoatem numai referinta catre antetul de prima pagina. Subsolul si `titlePg`
        # raman, deci numerotarea si asezarea in pagina nu se schimba.
        for s in doc.sections:
            sect = s._sectPr
            for ref in list(sect.findall(qn("w:headerReference"))):
                sect.remove(ref)
    return doc


def din_parametri(ref):
    """Formatul Scolii Doctorale UB Drept, singurul fara sablon propriu."""
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(ref["margini_cm"]["sus"])
        s.bottom_margin = Cm(ref["margini_cm"]["jos"])
        s.left_margin = Cm(ref["margini_cm"]["stanga"])
        s.right_margin = Cm(ref["margini_cm"]["dreapta"])

    normal = doc.styles["Normal"]
    normal.font.name = ref["font"]
    normal.font.size = Pt(ref["corp_pt"])
    _font_complet(normal.element, ref["font"])
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(ref["alineat_cm"])
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for nivel in (1, 2, 3):
        try:
            stil = doc.styles[f"Heading {nivel}"]
        except KeyError:
            continue
        stil.font.name = ref["font"]
        stil.font.size = Pt(ref["titlu_pt"] - nivel + 1)
        stil.font.bold = True
        stil.font.color.rgb = None
        _font_complet(stil.element, ref["font"])
    return doc


def rezumat(doc, cale: Path, academic: bool):
    sect = doc.sections[0]._sectPr
    anteturi = sorted(x for x in
                      (r.get(qn("w:type")) for r in sect.findall(qn("w:headerReference"))) if x)
    subsoluri = sorted(x for x in
                       (r.get(qn("w:type")) for r in sect.findall(qn("w:footerReference"))) if x)
    prima_diferita = sect.find(qn("w:titlePg")) is not None
    print(f"{cale}")
    print(f"  profil: {'academic UB Drept' if academic else 'stil de casa, din sablon'}")
    print(f"  antet: {anteturi or 'niciunul'}"
          f"{', doar pe prima pagina' if prima_diferita and anteturi else ''}")
    print(f"  subsol: {subsoluri or 'niciunul'}")


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Document nou in stilul de casa")
    ap.add_argument("--output", required=True, type=Path)
    ap.add_argument("--titlu", help="titlul actului, scris ca Heading 1")
    ap.add_argument("--fara-antet", action="store_true",
                    help="scoate logo-ul, pentru documente interne")
    ap.add_argument("--academic", action="store_true",
                    help="format Scoala Doctorala UB Drept, construit din parametri")
    ap.add_argument("--autor", default="Adrian Zamfir")
    args = ap.parse_args(argv)

    doc = din_parametri(ACADEMIC) if args.academic else din_sablon(args.fara_antet)

    if args.titlu:
        doc.add_heading(args.titlu, level=1)

    core = doc.core_properties
    core.author = args.autor
    core.last_modified_by = args.autor

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))
    # Word refuza un pachet in care Ignorable enumera prefixe nedeclarate, chiar daca
    # fisierul e XML valid si se deschide fara reproa in python-docx.
    repara_pachet.asigura(args.output)
    rezumat(doc, args.output, args.academic)
    return 0


if __name__ == "__main__":
    sys.exit(main())
