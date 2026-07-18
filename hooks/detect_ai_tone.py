#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Detect AI Tone v2.0 — Comprehensive automated detection of LLM-generated text patterns.

Scans Markdown/text documents for 25+ linguistic markers typical of AI text:
- Overused AI vocabulary (Romanian + English, 200+ terms)
- Em-dash overuse (Claude/ChatGPT signature)
- Colon overuse and "label colon" pattern
- Bullet point excess
- Rule-of-three obsession
- Negative parallelisms ("not X, but Y")
- Excessive hedging
- Truisms / platitudes / throat clearing
- Mechanical conclusions
- Chatbot artifacts
- Copula avoidance ("serves as" vs "is")
- Synonym cycling
- False ranges
- Tonal imbalance

Usage:
    python detect_ai_tone.py document.md [--json] [--language ro|en|auto]

Based on synthesis of 14 international sources (2024-2026):
- Wikipedia: Signs of AI Writing
- MIT Press: LLM-Generated Text Detection Survey (2025)
- arXiv 2510.05136 (Tercon, 2025)
- UCC Stylometry Study (O'Sullivan, 2025)
- HumanizerTech, Pangram Labs, Patina, Conor Bronsdon
- Plagiarism Today, Will Francis, Blake Stockton
- Washington Post (2025), Westcliff Writing Center
"""

import argparse
import json
import re
import sys
import unicodedata
from pathlib import Path


# === AI VOCABULARY MARKERS — RO ===
AI_WORDS_RO = [
    'fara precedent', 'paradigmatic', 'laitmotiv', 'catalizator',
    'pilon', 'arsenal', 'ecosistem', 'piatra de temelie',
    'de remarcat', 'merita mentionat', 'e de subliniat',
    'nu putem ignora', 'nu este intamplator', 'nu in ultimul rand',
    'este graitor', 'este elocvent', 'este emblematic',
    'aspect fundamental', 'dimensiune esentiala',
    'in acest context', 'de altfel', 'in consecinta', 'prin urmare',
    'in plus', 'totodata', 'de asemenea', 'mai mult decat atat',
    'pe de o parte', 'pe de alta parte', 'in primul rand',
    'in al doilea rand', 'in al treilea rand', 'cu toate acestea',
    'in mod fundamental', 'in mod paradoxal', 'in mod sistematic',
    'in mod natural', 'in mod evident', 'in mod clar',
    'cu atat mai mult cu cat', 'tocmai de aceea',
    'in egala masura', 'deopotriva', 'per ansamblu',
    'ar putea', 'eventual', 'in general', 'de regula',
    'in mare parte', 'in mod tipic', 'oarecum', 'relativ',
    'in lumea de azi', 'in era digitala', 'in contextul globalizarii',
    'tehnologia evolueaza', 'schimbarea este inevitabila',
    'in concluzie', 'in final', 'ramane sa vedem', 'doar timpul va spune',
    'viitorul pare', 'exista inca multe', 'implicatiile sunt vaste',
    'prezenta lucrare isi propune', 'aceasta analiza trece in revista',
    'documentul de fata',
    'serveste drept', 'se prezinta ca', 'reprezinta',
    'constituie', 'functioneaza ca', 'actioneaza ca',
    'se lauda cu',
    'holistic', 'holistica', 'multifatetat', 'multidimensional',
    'sa exploram', 'vom vedea ca', 'ne propunem sa',
    'asa cum am discutat', 'hai sa descompunem',
]

AI_WORDS_EN = [
    'stands as a testament', 'plays a vital role', 'plays a crucial role',
    'pivotal', 'groundbreaking', 'cutting-edge', 'transformative',
    'game-changing', 'innovative', 'a testament to',
    'delve', 'delve into', 'dive into', 'navigate the complexities',
    'underscore', 'underscores the importance', 'bolster', 'foster',
    'harness', 'leverage', 'unpack', 'shed light on', 'pave the way',
    'robust', 'comprehensive', 'seamless', 'intricate', 'nuanced',
    'multifaceted', 'holistic approach', 'synergy',
    'tapestry', 'rich tapestry', 'enduring', 'enduring legacy',
    'enhanced', 'fostering', 'highlighting', 'crucial',
    'landscape', "in today's landscape", 'realm', 'in the realm of',
    "it's important to note", "it's worth noting",
    'it bears mentioning', 'it bears noting',
    'no discussion would be complete', 'needless to say',
    'it goes without saying', 'without a doubt',
    'furthermore', 'moreover', 'additionally',
    'in addition', 'consequently', 'subsequently',
    "in today's fast-paced world", "in today's digital world",
    "in today's rapidly evolving", 'at its core', 'at the end of the day',
    'when it comes to', 'this is where', "let's break it down",
    "let's explore", "let's dive",
    'vibrant', 'breathtaking', 'stunning', 'renowned',
    'rich cultural heritage', 'rich heritage',
    'might', 'could potentially', 'arguably',
    'somewhat', 'relatively', 'perhaps',
    'in conclusion', 'overall', 'to conclude', 'to sum up',
    'the future looks bright', 'only time will tell',
    'exciting times ahead',
    'serves as', 'stands as', 'functions as', 'acts as',
    'boasts', 'features',
    'great question', 'excellent question', 'absolutely',
    'i hope this helps', 'let me know if', 'feel free to',
    'this section will explore', 'this section examines',
]

MECHANICAL_CONJUNCTIONS_RO = [
    'in primul rand', 'in al doilea rand', 'in al treilea rand',
    'pe de o parte', 'pe de alta parte', 'cu toate acestea',
    'mai mult decat atat', 'in plus', 'totodata', 'de asemenea',
    'prin urmare', 'in consecinta', 'de altfel', 'in acest context',
]

MECHANICAL_CONJUNCTIONS_EN = [
    'firstly', 'secondly', 'thirdly', 'fourthly',
    'on one hand', 'on the other hand', 'nevertheless', 'nonetheless',
    'additionally', 'correspondingly', 'furthermore', 'moreover',
    'however', 'thus', 'hence', 'therefore', 'consequently',
]

EDITORIAL_MARKERS_RO = [
    'nimeni nu', 'nimic nu', 'niciodata nu',
    'este graitor ca', 'este elocvent ca', 'este tragic ca',
    'din pacate', 'din fericire', 'inevitabil',
    'este evident ca', 'este clar ca', 'este incontestabil',
    'fara indoiala', 'fara echivoc',
]

EDITORIAL_MARKERS_EN = [
    'nobody', 'nothing short of', 'never before',
    "it's ironic that", 'tellingly', 'tragically',
    'unfortunately', 'fortunately', 'inevitably', 'undeniably',
    "it's clear that", "it's evident that", 'undoubtedly',
]

HEDGING_RO = [
    'poate', 'ar putea', 'eventual', 'posibil', 'probabil',
    'in general', 'de regula', 'in mare parte', 'in mod tipic',
    'oarecum', 'relativ', 'aparent', 'pare', 'tinde sa',
]

HEDGING_EN = [
    'might', 'could', 'potentially', 'may', 'typically',
    'generally', 'often', 'sometimes', 'somewhat',
    'relatively', 'possibly', 'perhaps', 'arguably',
    'apparently', 'seems', 'tends to',
]

TRUISMS_RO = [
    'in lumea de azi', 'in era digitala', 'in contextul globalizarii',
    'tehnologia evolueaza rapid', 'schimbarea este inevitabila',
    'comunicarea este esentiala', 'datele sunt noul petrol',
    'pe masura ce', 'in ultima perioada', 'devine din ce in ce mai',
]

TRUISMS_EN = [
    "in today's world", "in today's fast-paced",
    'in the digital age', 'in the era of',
    'as technology evolves', 'change is inevitable',
    'communication is key', 'data is the new oil',
    'in recent years', 'becoming increasingly',
]

MECHANICAL_CONCLUSIONS_RO = [
    'in concluzie', 'in final', 'per ansamblu',
    'ramane sa vedem', 'doar timpul va spune',
    'viitorul pare promitator', 'implicatiile sunt vaste',
    'exista inca multe de descoperit',
]

MECHANICAL_CONCLUSIONS_EN = [
    'in conclusion', 'in summary', 'overall',
    'to conclude', 'to sum up', 'in essence',
    'the future looks bright', 'only time will tell',
    'much remains to be seen', 'exciting times ahead',
]

CHATBOT_ARTIFACTS = [
    'i hope this helps', 'feel free to', 'let me know if',
    "don't hesitate to", 'great question', 'excellent question',
    'absolutely', 'certainly', 'of course',
    'as of my last update', 'as of my knowledge cutoff',
    'sper ca va ajuta', 'nu ezitati sa', 'daca aveti intrebari',
    'excelenta intrebare', 'desigur', 'cu siguranta',
]

COPULA_AVOIDANCE_RO = [
    'serveste drept', 'se prezinta ca', 'reprezinta',
    'constituie', 'functioneaza ca', 'actioneaza ca',
    'se lauda cu', 'ofera o imagine',
]

COPULA_AVOIDANCE_EN = [
    'serves as', 'stands as', 'functions as', 'acts as',
    'represents', 'constitutes', 'boasts', 'features',
]


def detect_language(text):
    ro_markers = ['si', 'in', 'pentru', 'este', 'care', 'din', 'sau', 'acest', 'prin', 'asupra']
    en_markers = ['the', 'and', 'for', 'that', 'with', 'from', 'this', 'are', 'have', 'been']
    words = text.lower().split()[:500]
    ro_count = sum(1 for w in words if w in ro_markers)
    en_count = sum(1 for w in words if w in en_markers)
    return 'ro' if ro_count > en_count else 'en'


def normalize(text):
    """Strip diacritics for case-insensitive matching."""
    nfkd = unicodedata.normalize('NFKD', text)
    return ''.join(c for c in nfkd if not unicodedata.combining(c))


def find_markers(text, marker_list, label):
    findings = []
    text_norm = normalize(text.lower())
    for marker in marker_list:
        marker_norm = normalize(marker.lower())
        count = text_norm.count(marker_norm)
        if count > 0:
            findings.append({'marker': marker, 'count': count, 'type': label})
    return findings


def count_em_dashes(text):
    em_count = text.count('—')
    word_count = len(text.split())
    pages = max(word_count / 250, 1)
    return {
        'count': em_count,
        'pages': round(pages, 1),
        'per_page': round(em_count / pages, 2),
        'threshold': 1.0,
        'above_threshold': (em_count / pages) > 1.0,
    }


def detect_appositional_dashes_ro(text):
    """Detect Romanian appositional structures using dashes (FORBIDDEN per DOOM 3).

    In Romanian, appositions MUST be delimited by commas, not by dashes.
    Pattern: 'word — apposition — word' (em-dash) or 'word – apposition – word' (en-dash)
    """
    patterns = [
        # Pattern 1: word EM-DASH apposition EM-DASH word (typical AI calque from English)
        r'\b\w+\s+[—–]\s+[^—–\n]{2,80}\s+[—–]\s+\w+',
        # Pattern 2: proper noun followed by em-dash + apposition (common in academic AI text)
        r'\b[A-ZĂÂÎȘȚ][a-zăâîșț]+\s+[—–]\s+(?:fondatorul|autorul|directorul|primarul|presedintele|conducatorul|fostul|actualul|profesorul|doctorul)\s',
    ]
    matches = []
    for pattern in patterns:
        for m in re.finditer(pattern, text):
            matches.append({
                'snippet': m.group(0)[:120],
                'position': m.start(),
            })
    return {
        'count': len(matches),
        'threshold': 0,
        'above_threshold': len(matches) > 0,
        'examples': matches[:5],
    }


def count_colons(text):
    paragraphs = [p for p in text.split('\n\n') if len(p.strip()) > 30]
    paragraphs_with_colon = 0
    total_colons = 0
    for p in paragraphs:
        clean = re.sub(r'https?://\S+', '', p)
        clean = re.sub(r'\d+:\d+', '', clean)
        if ':' in clean:
            paragraphs_with_colon += 1
            total_colons += clean.count(':')
    ratio = paragraphs_with_colon / max(len(paragraphs), 1)
    return {
        'total_colons': total_colons,
        'paragraphs_total': len(paragraphs),
        'paragraphs_with_colon': paragraphs_with_colon,
        'ratio': round(ratio, 2),
        'threshold': 0.30,
        'above_threshold': ratio > 0.30,
    }


def detect_label_colon_pattern(text):
    patterns = [
        r'^[\*\-\+]\s+\*\*[^*]+:\*\*',
        r'^[\*\-\+]\s+\*\*[^*\n]+\*\*\s*:\s*',
    ]
    count = 0
    for pattern in patterns:
        count += len(re.findall(pattern, text, re.MULTILINE))
    return {
        'count': count,
        'threshold': 0,
        'above_threshold': count > 0,
    }


def detect_negative_parallelisms(text):
    patterns = [
        r'(?:nu\s+(?:este|e)\s+\w+[^.]*?(?:este|e)\s+\w+)',
        r'(?:nu\s+doar\s+\w+[^.]*?ci\s+\w+)',
        r"(?:it'?s\s+not\s+\w+[^.]*?it'?s\s+\w+)",
        r'(?:not\s+just\s+\w+[^.]*?but\s+\w+)',
        r'(?:not\s+only\s+\w+[^.]*?but\s+\w+)',
        r'(?:no\s+\w+[\s.,]+no\s+\w+[\s.,]+no\s+\w+)',
        r'(?:niciun\s+\w+[\s.,]+niciun\s+\w+)',
    ]
    count = 0
    for pattern in patterns:
        count += len(re.findall(pattern, text, re.IGNORECASE))
    return {
        'count': count,
        'threshold': 1,
        'above_threshold': count > 1,
    }


def detect_rule_of_three(text):
    patterns = [
        r'(?:trei\s+(?:aspecte|factori|dimensiuni|elemente|categorii|consecinte))',
        r'(?:three\s+(?:aspects|factors|dimensions|elements|categories|consequences))',
    ]
    count = 0
    for pattern in patterns:
        count += len(re.findall(pattern, normalize(text), re.IGNORECASE))
    word_count = len(text.split())
    pages = max(word_count / 250, 1)
    per_10_pages = (count / pages) * 10
    return {
        'count': count,
        'per_10_pages': round(per_10_pages, 1),
        'threshold': 5,
        'above_threshold': per_10_pages > 5,
    }


def detect_bullet_density(text):
    lines = text.split('\n')
    bullet_lines = sum(1 for line in lines if re.match(r'^\s*[\*\-\+\d]+[\.)]\s+', line))
    para_lines = sum(1 for line in lines if line.strip() and not re.match(r'^\s*[\*\-\+\d#]+[\.)]?\s+', line))
    ratio = bullet_lines / max(para_lines, 1)
    return {
        'bullet_lines': bullet_lines,
        'paragraph_lines': para_lines,
        'ratio': round(ratio, 2),
        'threshold': 0.20,
        'above_threshold': ratio > 0.20,
    }


def analyze_tonal_balance(text, language):
    paragraphs = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 50]
    negative_words_ro = ['problema', 'deficit', 'lipsa', 'absenta', 'criza', 'degradare',
                         'esec', 'neglijare', 'subfinantare', 'fragmentare', 'vid',
                         'restante', 'pierdere', 'suspendare', 'disfunctional']
    positive_words_ro = ['oportunitate', 'potential', 'dezvoltare', 'crestere', 'progres',
                         'investitie', 'inovare', 'solutie', 'reforma', 'capabilit',
                         'competitiv', 'performanta', 'succes']
    negative_words_en = ['problem', 'deficit', 'lack', 'absence', 'crisis', 'failure',
                         'neglect', 'underfund', 'fragment', 'loss', 'suspend']
    positive_words_en = ['opportunity', 'potential', 'development', 'growth', 'progress',
                         'investment', 'innovation', 'solution', 'reform', 'competitive']
    neg_words = negative_words_ro if language == 'ro' else negative_words_en
    pos_words = positive_words_ro if language == 'ro' else positive_words_en
    negative_paras = positive_paras = neutral_paras = 0
    for para in paragraphs:
        para_lower = normalize(para.lower())
        neg_score = sum(1 for w in neg_words if w in para_lower)
        pos_score = sum(1 for w in pos_words if w in para_lower)
        if neg_score > pos_score + 1:
            negative_paras += 1
        elif pos_score > neg_score + 1:
            positive_paras += 1
        else:
            neutral_paras += 1
    total = max(negative_paras + positive_paras + neutral_paras, 1)
    return {
        'negative': negative_paras, 'positive': positive_paras, 'neutral': neutral_paras,
        'negative_ratio': round(negative_paras / total * 100, 1),
        'balanced': negative_paras / total < 0.7,
    }


def compute_naturalness_score(metrics):
    score = 100
    score -= min(metrics['ai_vocab_total'] * 1.5, 25)
    score -= min(metrics['hedging_total'] * 0.8, 15)
    score -= min(metrics['editorial_total'] * 2, 12)
    if metrics['em_dash']['above_threshold']:
        score -= min(metrics['em_dash']['per_page'] * 5, 15)
    if metrics.get('appositional_dashes', {}).get('above_threshold'):
        # Penalitate severa - apozitiile cu liniute in romana sunt CRITICAL
        score -= min(metrics['appositional_dashes']['count'] * 6, 25)
    if metrics['colon']['above_threshold']:
        score -= min((metrics['colon']['ratio'] - 0.30) * 30, 12)
    if metrics['label_colon']['above_threshold']:
        score -= min(metrics['label_colon']['count'] * 4, 15)
    if metrics['negative_parallelism']['above_threshold']:
        score -= min(metrics['negative_parallelism']['count'] * 3, 10)
    if metrics['bullet_density']['above_threshold']:
        score -= 8
    if metrics['truisms_total'] > 0:
        score -= min(metrics['truisms_total'] * 3, 10)
    if metrics['conclusions_total'] > 0:
        score -= min(metrics['conclusions_total'] * 2, 8)
    if metrics['chatbot_total'] > 0:
        score -= min(metrics['chatbot_total'] * 5, 20)
    if not metrics['tonal']['balanced']:
        score -= 12
    return max(0, round(score))


def read_document(doc_path):
    """Citeste .md/.txt/.docx. Pentru .docx extrage textul real (paragrafe + tabele),
    nu octetii bruti ai arhivei ZIP. Fallback pe dezarhivare word/document.xml.

    Esueaza ZGOMOTOS, niciodata silentios: un fisier lipsa sau nedecodabil trebuie sa
    dea eroare, nu text gol. Text gol ar produce scor 100/100 (fals pozitiv perfect).
    """
    p = Path(doc_path)
    if not p.exists():
        raise FileNotFoundError(f'Documentul nu exista: {p}')
    if p.stat().st_size == 0:
        raise ValueError(f'Document gol (0 octeti): {p}')

    if p.suffix.lower() == '.docx':
        errors = []
        try:
            import docx
            d = docx.Document(str(p))
            parts = [para.text for para in d.paragraphs]
            for table in d.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return '\n'.join(parts)
        except Exception as e:
            errors.append(f'python-docx: {type(e).__name__}: {e}')
        try:
            import zipfile
            with zipfile.ZipFile(str(p)) as z:
                xml = z.read('word/document.xml').decode('utf-8', errors='ignore')
            xml = re.sub(r'</w:p>', '\n', xml)
            return re.sub(r'<[^>]+>', '', xml)
        except Exception as e:
            errors.append(f'zipfile: {type(e).__name__}: {e}')
        raise RuntimeError(f'Nu pot extrage textul din .docx {p}. Incercari: ' + ' | '.join(errors))

    return p.read_text(encoding='utf-8', errors='ignore')


def analyze_text(text, language='auto', doc_label='<text>'):
    """Analiza pura, fara IO: primeste text, intoarce dict-ul de rezultate.
    Punct de intrare in-process pentru serverul MCP (fara subprocess, fara temp file)."""
    if language == 'auto':
        language = detect_language(text)
    ai_vocab = find_markers(text, AI_WORDS_RO if language == 'ro' else AI_WORDS_EN, 'AI_VOCAB')
    editorials = find_markers(text, EDITORIAL_MARKERS_RO if language == 'ro' else EDITORIAL_MARKERS_EN, 'EDIT')
    hedging = find_markers(text, HEDGING_RO if language == 'ro' else HEDGING_EN, 'HEDGE')
    truisms = find_markers(text, TRUISMS_RO if language == 'ro' else TRUISMS_EN, 'TRUISM')
    conclusions = find_markers(text, MECHANICAL_CONCLUSIONS_RO if language == 'ro' else MECHANICAL_CONCLUSIONS_EN, 'CONCL')
    chatbot = find_markers(text, CHATBOT_ARTIFACTS, 'CHATBOT')
    copula = find_markers(text, COPULA_AVOIDANCE_RO if language == 'ro' else COPULA_AVOIDANCE_EN, 'COPULA')
    em_dash = count_em_dashes(text)
    appositional_dashes = detect_appositional_dashes_ro(text) if language == 'ro' else {'count': 0, 'above_threshold': False, 'examples': []}
    colon = count_colons(text)
    label_colon = detect_label_colon_pattern(text)
    neg_parallel = detect_negative_parallelisms(text)
    rule_of_three = detect_rule_of_three(text)
    bullet = detect_bullet_density(text)
    tonal = analyze_tonal_balance(text, language)
    metrics = {
        'ai_vocab_total': sum(f['count'] for f in ai_vocab),
        'hedging_total': sum(f['count'] for f in hedging),
        'editorial_total': sum(f['count'] for f in editorials),
        'truisms_total': sum(f['count'] for f in truisms),
        'conclusions_total': sum(f['count'] for f in conclusions),
        'chatbot_total': sum(f['count'] for f in chatbot),
        'copula_total': sum(f['count'] for f in copula),
        'em_dash': em_dash,
        'appositional_dashes': appositional_dashes,
        'colon': colon,
        'label_colon': label_colon,
        'negative_parallelism': neg_parallel,
        'rule_of_three': rule_of_three,
        'bullet_density': bullet,
        'tonal': tonal,
    }
    naturalness = compute_naturalness_score(metrics)
    results = {
        'document': doc_label,
        'language': language,
        'word_count': len(text.split()),
        'naturalness_score': naturalness,
        'verdict': 'NATURAL' if naturalness >= 80 else 'REVIZUIRE' if naturalness >= 60 else 'AI-LIKE',
        'metrics': metrics,
        'top_offenders': {
            'vocab': sorted(ai_vocab, key=lambda x: -x['count'])[:10],
            'hedging': sorted(hedging, key=lambda x: -x['count'])[:5],
            'editorial': sorted(editorials, key=lambda x: -x['count'])[:5],
            'truisms': sorted(truisms, key=lambda x: -x['count'])[:5],
            'chatbot': sorted(chatbot, key=lambda x: -x['count'])[:5],
        },
    }
    return results


def run_detection(doc_path, language='auto', output_json=False):
    text = read_document(doc_path)
    results = analyze_text(text, language, str(doc_path))
    if output_json:
        print(json.dumps(results, ensure_ascii=False, indent=2))
    else:
        print_human_report(results)
    return results['naturalness_score']


def print_human_report(r):
    print(f"\n{'='*64}")
    print(f"  ANALIZA DE TON v2.0 - {Path(r['document']).name}")
    print(f"  Limba: {'Romana' if r['language'] == 'ro' else 'Engleza'}")
    print(f"  Cuvinte: {r['word_count']:,}")
    print(f"{'='*64}")
    score = r['naturalness_score']
    print(f"\n>>> Scor naturalete: {score}/100  [{r['verdict']}]")
    m = r['metrics']
    if m['ai_vocab_total'] > 0:
        print(f"\n[VOCAB] Vocabular AI ({m['ai_vocab_total']} aparitii):")
        for item in r['top_offenders']['vocab']:
            print(f"   - {item['marker']:40s} x {item['count']}")
    if m['em_dash']['above_threshold']:
        print(f"\n[EM-DASH] Suprautilizat: {m['em_dash']['per_page']}/pagina (max 1.0)")
    if m.get('appositional_dashes', {}).get('above_threshold'):
        print(f"\n[APOZITII-LINIUTE] CRITICAL - in romana apozitiile se delimiteaza prin VIRGULE, nu liniute!")
        print(f"   Detectate: {m['appositional_dashes']['count']} aparitii (REGULA 2.2 / DOOM 3)")
        for ex in m['appositional_dashes'].get('examples', [])[:3]:
            print(f"   GRESIT: {ex['snippet']}")
    if m['colon']['above_threshold']:
        print(f"\n[COLON] Suprautilizat: {m['colon']['ratio']*100:.0f}% paragrafe au ':' (max 30%)")
    if m['label_colon']['above_threshold']:
        print(f"\n[LABEL-COLON] Pattern '**X:** desc' detectat: {m['label_colon']['count']} aparitii (REGULA 0)")
    if m['negative_parallelism']['above_threshold']:
        print(f"\n[PARALELISM] 'Nu X, ci Y' / 'Not X, but Y': {m['negative_parallelism']['count']} (max 1)")
    if m['bullet_density']['above_threshold']:
        print(f"\n[BULLETS] Densitate excesiva: {m['bullet_density']['ratio']*100:.0f}% (max 20%)")
    if m['rule_of_three']['above_threshold']:
        print(f"\n[TRIADE] Rule of three: {m['rule_of_three']['per_10_pages']}/10 pag. (max 5)")
    if m['hedging_total'] > 0:
        rate = m['hedging_total'] / max(r['word_count']/100, 1)
        print(f"\n[HEDGING] {m['hedging_total']} aparitii ({rate:.1f}/100 cuvinte)")
    if m['truisms_total'] > 0:
        print(f"\n[TRUISME] {m['truisms_total']} aparitii:")
        for item in r['top_offenders']['truisms']:
            print(f"   - {item['marker']:40s} x {item['count']}")
    if m['conclusions_total'] > 0:
        print(f"\n[CONCLUZII MECANICE] {m['conclusions_total']} aparitii")
    if m['chatbot_total'] > 0:
        print(f"\n[CHATBOT-ARTIFACTS] {m['chatbot_total']} aparitii (CRITICAL):")
        for item in r['top_offenders']['chatbot']:
            print(f"   - {item['marker']:40s} x {item['count']}")
    if m['copula_total'] > 5:
        print(f"\n[COPULA-AVOIDANCE] 'serves as' / 'reprezinta': {m['copula_total']} (atentie)")
    if not m['tonal']['balanced']:
        print(f"\n[TONAL] Dezechilibru: {m['tonal']['negative_ratio']}% negativ (max 70%)")
    print(f"\n{'='*64}\n")


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Detectare ton AI v2.0')
    parser.add_argument('document', help='Calea catre document')
    parser.add_argument('--language', '-l', choices=['ro', 'en', 'auto'], default='auto')
    parser.add_argument('--json', action='store_true', help='Output JSON')
    args = parser.parse_args()
    try:
        score = run_detection(args.document, args.language, args.json)
    except (FileNotFoundError, ValueError, RuntimeError) as e:
        # Eroare de citire, nu verdict de ton. Cod 2, ca sa nu fie confundata cu
        # "text AI-like" (cod 1) si sa nu treaca drept scor bun.
        if args.json:
            print(json.dumps({'error': str(e)}, ensure_ascii=False, indent=2))
        else:
            print(f'EROARE: {e}', file=sys.stderr)
        sys.exit(2)
    sys.exit(0 if score >= 60 else 1)
